import streamlit as st
import openpyxl
import pandas as pd
from docx import Document

# Fungsi untuk menghitung kolom nilai dan keterangan
def calculate_rating(row):
    if pd.notnull(row['TOT POINT']):
        if row['TOT POINT'] <= 40:
            return 'E'
        elif row['TOT POINT'] <= 60:
            return 'D'
        elif row['TOT POINT'] <= 70:
            return 'C'
        elif row['TOT POINT'] <= 89:
            return 'B'
        else:
            return 'A'
    return ''

def calculate_keterangan(row):
    if row['nilai'] == 'E':
        return 'Sangat Buruk / tidak dilanjutkan kerjasamanya'
    elif row['nilai'] == 'D':
        return 'Buruk / dipertimbangkan kerjasamanya dan diberikan surat peringatan'
    elif row['nilai'] == 'C':
        return 'Cukup / dilanjutkan kerjasamanya'
    elif row['nilai'] == 'B':
        return 'Baik / dilanjutkan kerjasamanya'
    elif row['nilai'] == 'A':
        return 'Sangat Baik / dilanjutkan kerjasamanya'
    return ''

def read_excel_and_drop_nan(file_path, sheet_name):
    # Membaca file Excel
    wb = openpyxl.load_workbook(file_path)

    try:
        # Memilih sheet berdasarkan nama
        sheet = wb[sheet_name]
    except KeyError:
        print(f"Sheet dengan nama '{sheet_name}' tidak ditemukan.")
        return None
    else:
        # Mencari baris indeks yang berisi teks "NO Vendor_name, TOT.PO, RP/PP, KUALITAS, K3, L, TOT, POINT, Nilai, KETERANGAN"
        target_row_index = None
        for row in sheet.iter_rows(min_row=1, max_row=10):
            row_values = [cell.value for cell in row]
            if "NO" in row_values and "Vendor_name" in row_values and "TOT.PO" in row_values:
                target_row_index = row[0].row
                break

        if target_row_index is None:
            print("Tidak dapat menemukan baris yang sesuai di sheet", sheet_name)
            return None
        else:
            # Membaca DataFrame dari baris yang sesuai
            df = pd.read_excel(file_path, sheet_name=sheet_name, skiprows=target_row_index - 1, usecols="A:J")

            def drop_nan_rows_from_top(df):
                # Mencari baris pertama dengan semua nilai NaN
                first_nan_row_index = df.apply(lambda row: row.isnull().all(), axis=1).idxmax()

                # Menghapus baris-baris dengan semua nilai NaN dari bagian atas dataframe
                df = df.loc[first_nan_row_index:].copy()

                return df

            # Menghapus baris dengan nilai NaN dari bagian atas dataframe
            df_cleaned = df.drop(drop_nan_rows_from_top(df).index)

            # Menambahkan kolom 'bulan' berdasarkan kata ke-2 dari sheet_name
            df_cleaned['BULAN'] = sheet_name.split()[1]

            # Menambahkan kolom 'tahun' berdasarkan kata ke-3 dari sheet_name
            df_cleaned['TAHUN'] = sheet_name.split()[2]

            return df_cleaned

def read_all_sheets_in_excel(file_path):
    # Membaca file Excel
    wb = openpyxl.load_workbook(file_path)

    # Menginisialisasi list untuk menyimpan DataFrame dari setiap sheet
    all_dataframes = []

    for sheet_name in wb.sheetnames:
        df = read_excel_and_drop_nan(file_path, sheet_name)
        if df is not None:
            all_dataframes.append(df)

    # Menggabungkan semua DataFrame menjadi satu DataFrame
    result_df = pd.concat(all_dataframes, ignore_index=True)

    return result_df

# Fungsi untuk menampilkan DataFrame yang telah difilter berdasarkan vendor, rentang bulan, dan tahun
def display_filtered_data(df):
    # Tampilkan widget untuk memilih vendor
    selected_vendors = st.multiselect('Select Vendor:', df['Vendor_name'].unique())

    # Filter DataFrame berdasarkan vendor yang dipilih
    filtered_data = df[df['Vendor_name'].isin(selected_vendors)] if selected_vendors else df

    # Tampilkan widget untuk memilih tahun
    selected_years = st.multiselect('Select Year:', df['TAHUN'].unique())

    # Filter DataFrame berdasarkan tahun yang dipilih
    filtered_data = filtered_data[filtered_data['TAHUN'].isin(selected_years)] if selected_years else filtered_data

    # Tampilkan beberapa pilihan bulan
    months = df['BULAN'].unique()
    selected_months = st.multiselect('Select Months:', months)

    # Filter DataFrame berdasarkan bulan yang dipilih
    filtered_data = filtered_data[filtered_data['BULAN'].isin(selected_months)] if selected_months else filtered_data
    
    # Tampilkan DataFrame yang sudah difilter
    st.dataframe(filtered_data)
    
    
    
    # Hitung sum dan average (mean) dari kolom yang dipilih
    if not filtered_data.empty:
        sum_selected_columns = filtered_data['TOT.PO'].sum()
        mean_selected_columns = filtered_data[['RP/PP', 'KUALITAS', 'K3', 'L', 'TOT POINT']].mean()


        summary_df = pd.DataFrame({
            'Sum of TOT.PO': [sum_selected_columns],
            'Mean of RP/PP': [mean_selected_columns['RP/PP']],
            'Mean of KUALITAS': [mean_selected_columns['KUALITAS']],
            'Mean of K3': [mean_selected_columns['K3']],
            'Mean of L': [mean_selected_columns['L']],
            'Mean of TOT POINT': [mean_selected_columns['TOT POINT']]
        })
        def assign_grade(score):
            if score <= 40:
                return 'E'
            elif score <= 60:
                return 'D'
            elif score <= 70:
                return 'C'
            elif score <= 89:
                return 'B'
            else:
                return 'A'

        summary_df['Nilai'] = summary_df['Mean of TOT POINT'].apply(assign_grade)
        
        def assign_grade(score):
            if score <= 40:
                return 'Sangat Buruk / tidak dilanjutkan kerjasamanya'
            elif score <= 60:
                return 'Buruk / dipertimbangkan kerjasamanya dan diberikan surat peringatan'
            elif score <= 70:
                return 'Cukup / dilanjutkan kerjasamanya'
            elif score <= 89:
                return 'Baik / dilanjutkan kerjasamanya'
            else:
                return 'Sangat Baik / dilanjutkan kerjasamanya'

        summary_df['Keterangan'] = summary_df['Mean of TOT POINT'].apply(assign_grade)

        st.write("Summary:")
        st.dataframe(summary_df)
        return filtered_data, summary_df

# Fungsi untuk membuat dokumen Word berdasarkan template dan DataFrame
def generate_word_document(template_path, summary_df):
    # Baca template dokumen Word
    doc = Document(template_path)

    # Ganti placeholder dengan nilai dari DataFrame
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{Mean of RP/PP}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{Mean of RP/PP}', str(summary_df['Mean of RP/PP'][0]))
                    if '{Mean of KUALITAS}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{Mean of KUALITAS}', str(summary_df['Mean of KUALITAS'][0]))
                    if '{Mean of K3}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{Mean of K3}', str(summary_df['Mean of K3'][0]))
                    if '{Mean of L}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{Mean of L}', str(summary_df['Mean of L'][0]))
                    if '{Mean of TOT POINT}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{Mean of TOT POINT}', str(summary_df['Mean of TOT POINT'][0]))
                    if '{Vendor_name}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{Vendor_name}', str(filtered_data['Vendor_name'].iloc[0]))
                    if '{Keterangan}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{Keterangan}', str(summary_df['Keterangan'][0]))

    # Simpan dokumen Word
    output_path = 'summary_output.docx'
    doc.save(output_path)

    return output_path


# Load data from the uploaded file
uploaded_file = st.file_uploader("Upload your Excel file", type=["xlsx", "xls"])

if uploaded_file is not None:
    # Read all sheets in the Excel file and concatenate them into one DataFrame
    result_df = read_all_sheets_in_excel(uploaded_file)

    # Display filtered data based on user selection
    filtered_data, summary_df  = display_filtered_data(result_df)
    
    # Generate and download Word document
    template_path = "templates/Bismillah.docx"  # Path to your Word template
    word_output_path = generate_word_document(template_path, summary_df)
    st.download_button(label="Download Word Document", data=word_output_path, file_name="Bismillah.docx")
